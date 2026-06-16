import os
import io
import urllib.parse
from django.conf import settings
from pathlib import Path
import boto3
from botocore.exceptions import ClientError

# --- S3 Configuration ---
USE_S3 = os.environ.get('USE_S3', 'False').lower() == 'true'
AWS_ACCESS_KEY_ID = os.environ.get('AWS_ACCESS_KEY_ID')
AWS_SECRET_ACCESS_KEY = os.environ.get('AWS_SECRET_ACCESS_KEY')
AWS_STORAGE_BUCKET_NAME = os.environ.get('AWS_STORAGE_BUCKET_NAME')
AWS_REGION = os.environ.get('AWS_REGION', 'us-east-1')

def get_s3_client():
    if AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY:
        return boto3.client(
            's3',
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
            region_name=AWS_REGION
        )
    return boto3.client('s3', region_name=AWS_REGION)

def get_user_dir(user_id: int) -> Path:
    """Returns the Path directory for a specific user's files."""
    user_dir = Path(settings.MEDIA_ROOT) / "users" / str(user_id)
    os.makedirs(user_dir, exist_ok=True)
    return user_dir

# --- Word Document Binary Conversion Helpers ---

def convert_text_to_docx(text: str) -> bytes:
    """Creates a valid .docx binary from plain text."""
    import docx
    doc = docx.Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def convert_docx_to_text(file_bytes: bytes) -> str:
    """Extracts raw text from a .docx binary."""
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# --- PDF Document Binary Conversion Helpers ---

def convert_text_to_pdf(text: str) -> bytes:
    """Creates a valid .pdf binary from plain text using fpdf2."""
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, text=text)
    return bytes(pdf.output())

def convert_pdf_to_text(file_bytes: bytes) -> str:
    """Extracts raw text from a .pdf binary using PyMuPDF."""
    import fitz
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = []
        for page in doc:
            text.append(page.get_text())
        doc.close()
        return "\n".join(text)
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"

# ------------------------------------------------

def write_user_file(user_id: int, filename: str, content) -> int:
    """
    Writes file content (str or bytes) to user's folder.
    Automatically converts text to valid Word .docx binary if filename is .doc/.docx.
    Automatically converts text to valid PDF binary if filename is .pdf.
    Returns: file size in bytes.
    """
    is_word_doc = filename.lower().endswith(('.docx', '.doc'))
    is_pdf_doc = filename.lower().endswith('.pdf')

    if is_word_doc and isinstance(content, str):
        try:
            file_bytes = convert_text_to_docx(content)
        except Exception:
            file_bytes = content.encode('utf-8')
    elif is_pdf_doc and isinstance(content, str):
        try:
            file_bytes = convert_text_to_pdf(content)
        except Exception:
            file_bytes = content.encode('utf-8')
    elif isinstance(content, str):
        file_bytes = content.encode('utf-8')
    else:
        file_bytes = content

    # Write to local cache
    user_dir = get_user_dir(user_id)
    file_path = user_dir / filename
    with open(file_path, 'wb') as f:
        f.write(file_bytes)

    # Upload to S3 if enabled
    if USE_S3 and AWS_STORAGE_BUCKET_NAME:
        try:
            s3 = get_s3_client()
            s3_key = f"users/{user_id}/{filename}"
            s3.put_object(
                Bucket=AWS_STORAGE_BUCKET_NAME,
                Key=s3_key,
                Body=file_bytes
            )
        except Exception as e:
            print(f"Failed to upload to S3 for {filename}: {e}")

    return len(file_bytes)

def read_user_file(user_id: int, filename: str) -> bytes:
    """Reads file content from the user's folder (local cache or downloaded from S3)."""
    user_dir = get_user_dir(user_id)
    file_path = user_dir / filename

    # If S3 is enabled and file not in local cache, download from S3
    if USE_S3 and AWS_STORAGE_BUCKET_NAME and not file_path.exists():
        try:
            s3 = get_s3_client()
            s3_key = f"users/{user_id}/{filename}"
            response = s3.get_object(Bucket=AWS_STORAGE_BUCKET_NAME, Key=s3_key)
            file_bytes = response['Body'].read()
            # Cache it locally
            with open(file_path, 'wb') as f:
                f.write(file_bytes)
        except Exception as e:
            print(f"Failed to read from S3 for {filename}: {e}")
            # Fallback to reading local if download failed but file somehow exists
            if not file_path.exists():
                raise FileNotFoundError(f"File {filename} not found in S3 or local cache. Error: {e}")

    with open(file_path, 'rb') as f:
        return f.read()

def delete_user_file(user_id: int, filename: str):
    """Deletes a file from the user's folder (local cache and S3)."""
    user_dir = get_user_dir(user_id)
    file_path = user_dir / filename
    if file_path.exists():
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"Failed to delete local cache file {filename}: {e}")

    if USE_S3 and AWS_STORAGE_BUCKET_NAME:
        try:
            s3 = get_s3_client()
            s3_key = f"users/{user_id}/{filename}"
            s3.delete_object(Bucket=AWS_STORAGE_BUCKET_NAME, Key=s3_key)
        except Exception as e:
            print(f"Failed to delete {filename} from S3: {e}")

def list_user_files(user_id: int) -> list[str]:
    """Lists the filenames in the user's folder (from S3 if enabled, otherwise local cache)."""
    if USE_S3 and AWS_STORAGE_BUCKET_NAME:
        try:
            s3 = get_s3_client()
            prefix = f"users/{user_id}/"
            response = s3.list_objects_v2(Bucket=AWS_STORAGE_BUCKET_NAME, Prefix=prefix)
            filenames = []
            if 'Contents' in response:
                for obj in response['Contents']:
                    key = obj['Key']
                    filename = key[len(prefix):]
                    if filename:
                        filenames.append(filename)
            return filenames
        except Exception as e:
            print(f"Failed to list files from S3 for user {user_id}: {e}")

    # Fallback to local
    user_dir = get_user_dir(user_id)
    try:
        return [f for f in os.listdir(user_dir) if os.path.isfile(user_dir / f)]
    except FileNotFoundError:
        return []

def get_user_file_url(user_id: int, filename: str) -> str:
    """Returns the URL path for a user's file (S3 presigned URL if enabled, otherwise local media URL)."""
    if USE_S3 and AWS_STORAGE_BUCKET_NAME:
        try:
            s3 = get_s3_client()
            s3_key = f"users/{user_id}/{filename}"
            # Generate a secure presigned URL valid for 1 hour (3600 seconds)
            url = s3.generate_presigned_url(
                ClientMethod='get_object',
                Params={
                    'Bucket': AWS_STORAGE_BUCKET_NAME,
                    'Key': s3_key
                },
                ExpiresIn=3600
            )
            return url
        except Exception as e:
            print(f"Failed to generate presigned URL for {filename}: {e}")
            # Fallback to direct URL if presigned generation fails
            encoded_filename = urllib.parse.quote(filename)
            return f"https://{AWS_STORAGE_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/users/{user_id}/{encoded_filename}"

    encoded_filename = urllib.parse.quote(filename)
    return f"{settings.MEDIA_URL}users/{user_id}/{encoded_filename}"